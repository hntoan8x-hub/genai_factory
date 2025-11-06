from typing import Any, Dict
import logging
import asyncio
import os
# Giả định thư viện xử lý DOCX (ví dụ: python-docx)
try:
    import docx
except ImportError:
    docx = None

from shared_libs.ingestion.base.base_loader import BaseLoader
from shared_libs.utils.exceptions import GenAIFactoryError

logger = logging.getLogger(__name__)

class DocxLoader(BaseLoader):
    """
    A concrete loader for DOCX files. Implements asynchronous loading by offloading 
    the synchronous DOCX reading operation to a thread.
    """
    
    @property
    def supported_file_type(self) -> str:
        return "docx"

    def _extract_text_sync(self, file_path: str) -> Dict[str, Any]:
        """
        Internal synchronous method to read the DOCX file and extract text/metadata.
        (THIS IS THE BLOCKING OPERATION)
        """
        if docx is None:
            raise ImportError("The 'python-docx' library is required for DocxLoader.")
            
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        raw_text = []
        try:
            document = docx.Document(file_path)
            for paragraph in document.paragraphs:
                raw_text.append(paragraph.text)
            
            full_text = '\n'.join(raw_text)
            
            return {
                'content': full_text,
                'metadata': {
                    'source_path': file_path,
                    'file_type': self.supported_file_type,
                    'num_paragraphs': len(document.paragraphs)
                }
            }
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise GenAIFactoryError(f"DOCX extraction failed for {file_path}.")

    async def async_load(self, file_path_or_data: Any) -> Dict[str, Any]:
        """
        Hardening: Asynchronously loads the DOCX by running the synchronous logic 
        in a separate thread using asyncio.to_thread.
        """
        if not isinstance(file_path_or_data, str):
            raise TypeError("DocxLoader expects a file path string for loading.")
            
        logger.info(f"Starting async extraction for DOCX: {file_path_or_data}")
        
        # Offload blocking _extract_text_sync to a thread
        return await asyncio.to_thread(self._extract_text_sync, file_path_or_data)